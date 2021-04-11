import os
from docx import Document
from PyPDF2 import PdfFileMerger, PdfFileReader
import tkinter as tk

# Getting the current working directory path in original_path variable
original_path = os.getcwd()

def get_file_name(files_found, extension):
    """
    Function responsible for finding the path of the particular file.
    Input : 
    files_found(dictionary) - Dictionary of files found
    extension(string) - Extension of the file

    Output : 
    list - Returns the path of all the files
    """
    file_name = []
    for key, value in files_found.items():
        if key == extension:
            for val in value:
                if(val[1][-1] == '\\'):
                    file_name.append(val[1] + val[0] + '.' + extension)
                else:
                    file_name.append(
                        val[1] + '\\' + val[0] + '.' + extension)
    return file_name


def merge_txt(files_found, open_button, file_to_search):
    """
            Function is responsible for merging the .txt files.
            Input :
            files_found(dictionary) - Dictionary of files found
            open_button(Object) - Object of open_button related to txt files
            """
    file_name = get_file_name(files_found, 'txt')
    with open(original_path + '\\Merged_' + file_to_search.get() + '.txt', 'w') as outfile:
        for fname in file_name:
            with open(fname) as infile:
                data = infile.read()
                outfile.write(data)
    outfile.close()

    # Disabled state of button is changed here
    open_button['state'] = 'normal'


def merge_docx(files_found, open_button, file_to_search):
    """
            Function is responsible for merging the .docx files.
            Input :
            files_found(dictionary) - Dictionary of files found
            open_button(Object) - Object of open_button related to txt files
            """
    file_name = get_file_name(files_found, 'docx')
    final_document = Document()
    for file in file_name:
        source_document = Document(file)
        for paragraph in source_document.paragraphs:
            text = paragraph.text
            final_document.add_paragraph(text)
    final_document.save(original_path + '\\Merged_' +
                        file_to_search.get() + '.docx')

    # Disabled button is changed here
    open_button['state'] = 'normal'


def merge_pdf(files_found, open_button, file_to_search):
    """
            Function is responsible for merging the .pdf files.
            Input :
            files_found(dictionary) - Dictionary of files found
            open_button(Object) - Object of open_button related to txt files
            """
    file_name = get_file_name(files_found, 'pdf')
    merged_pdf = PdfFileMerger()
    for file in file_name:
        merged_pdf.append(PdfFileReader(file, 'rb'))

    merged_pdf.write(original_path + '\\Merged_' +
                     file_to_search.get() + '.pdf')

    # Disabled button is changed here
    open_button['state'] = 'normal'


def open_file(extension, file_to_search):
    """
            Function responsible for opening the merged file.
            Input :
            extension(string) - The file extension to open.
            """
    os.startfile(original_path + '\\Merged_' +
                      file_to_search.get() + '.' + extension)


def merge_func(merge_button_lis, open_button_lis, files_found, file_to_search):
    """
    Function responsible for adding event handlers to Merge and Open Buttons
    Input :
    merge_button_lis(list) - Contains the list of objects of merge buttons.
    open_button_lis(list) - Contains the list of objects of open buttons.
    files_found(Dictionary) - Dictionary of files found
    """
    for i in range(len(merge_button_lis)):
        if merge_button_lis[i]['text'].find('docx') >= 0:
            merge_button_lis[i].bind('<Button-1>', lambda event, files_found=files_found,
                                     open_button=open_button_lis[i]: merge_docx(files_found, open_button, file_to_search))
            open_button_lis[i].bind(
                '<Button-1>', lambda event, extension='docx': open_file(extension, file_to_search))

        elif merge_button_lis[i]['text'].find('pdf') >= 0:
            merge_button_lis[i].bind('<Button-1>', lambda event, files_found=files_found,
                                     open_button=open_button_lis[i]: merge_pdf(files_found, open_button, file_to_search))
            open_button_lis[i].bind(
                '<Button-1>', lambda event, extension='pdf': open_file(extension, file_to_search))

        elif merge_button_lis[i]['text'].find('txt') >= 0:
            merge_button_lis[i].bind('<Button-1>', lambda event, files_found=files_found,
                                     open_button=open_button_lis[i]: merge_txt(files_found, open_button, file_to_search))
            open_button_lis[i].bind(
                '<Button-1>', lambda event, extension='txt': open_file(extension, file_to_search))


def show_files(files_found, window):
    """
            Function responsible for displaying the Merge and Open Button as well as Labels on the UI
            Input : 
            files_found(Dictionary) - Dictionary of files found
            """
    x_button = 50
    y_button = 150
    x_label = 200
    y_label = 156
    x_open = 50
    y_open = 500
    merge_button_lis = []
    open_button_lis = []

    for key, value in files_found.items():
        if key == 'docx' or key == 'txt' or key == 'pdf':
            # Displaying Merge Button
            merge_button_lis.append(tk.Button(
                window, text='Merge ' + key + ' Files', fg='black', height=2, width=15, padx=2, pady=2))
            merge_button_lis[-1].place(x=x_button, y=y_button)
            y_button += 100

            # Displaying Labels
            label_lis = []
            total_files = str(len(value))
            label_lis.append(tk.Label(window, text='Total ' + key +
                                        ' files found = ' + total_files, font=('', 10)))
            label_lis[-1].place(x=x_label, y=y_label)
            y_label += 100

            # Displaying Open Button
            open_button_lis.append(tk.Button(window, text='Open ' + key + ' File',
                                               fg='black', height=2, width=16, padx=2, pady=2, state=tk.DISABLED))
            open_button_lis[-1].place(x=x_open, y=y_open)
            x_open += 250

    return merge_button_lis, open_button_lis
